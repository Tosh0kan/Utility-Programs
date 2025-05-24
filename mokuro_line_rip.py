import os
import json
import argparse
from docx import Document

def line_rip(dir_path) -> None:
    def json_rip() -> dict:
        dir_walk = os.walk(dir_path)
        pages = {}
        for root, dirs, files in dir_walk:
            for file in files:
                if 'json' in file:
                    with open(os.path.join(root, file), 'r', encoding='utf-8') as f:
                        data = json.loads(f.read())
                    all_lines = [n['lines'] for n in data['blocks']]
                    procced_lines = []
                    for line in all_lines:
                        if len(line) == 1:
                            procced_lines.append(line[0])
                        else:
                            for n in range(len(line)):
                                procced_lines.append(line[n])
                    pages.setdefault(file.split('.')[0], procced_lines)
                else:
                    pass
        return pages

    def txt_save(pages: dict) -> None:
        document = Document()
        for key, value in pages.items():
            document.add_heading('Page ' + key, level=2)

            for n in range(2):
                document.add_paragraph()

            for n in value:
                document.add_paragraph(n)

            for n in range(3):
                document.add_paragraph()
        document.save(dir_path + r'\speech_extracted.docx')

    txt_save(json_rip())


if __name__ == '__main__':
    parser = argparse.ArgumentParser()
    parser.add_argument('dir_path', type=str)
    args = parser.parse_args()
    line_rip(args.dir_path)
